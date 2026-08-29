import os
import re
import json
import zipfile
import docx
from xml.etree import ElementTree as ET

SOLVED_DIR = 'QUESTIONS/PAST YEAR /COMPLETE SOLVED'
OUT_IMG_DIR = 'public/docx_images'
OUT_JSON = 'src/data/questions.json'

os.makedirs(OUT_IMG_DIR, exist_ok=True)

SECTION_MAP = {
    'general aptitude': 'General Aptitude',
    'general aptitude (ga)': 'General Aptitude',
    'ga': 'General Aptitude',
    'engineering mathematics': 'Engineering Mathematics',
    '1 — engineering mathematics': 'Engineering Mathematics',
    'section 1: engineering mathematics': 'Engineering Mathematics',
    'farm power and machinery': 'Farm Power and Machinery',
    'farm machinery & power': 'Farm Power and Machinery',
    'farm machinery': 'Farm Power and Machinery',
    'farm power': 'Farm Power and Machinery',
    'soil and water conservation engineering': 'Soil and Water Conservation Engineering',
    'soil & water conservation engineering': 'Soil and Water Conservation Engineering',
    'soil and water engineering': 'Soil and Water Conservation Engineering',
    'agricultural process engineering': 'Agricultural Process Engineering',
    'agricultural processing engineering': 'Agricultural Process Engineering',
    'agricultural processing': 'Agricultural Process Engineering',
    'food process engineering': 'Agricultural Process Engineering',
    'dairy and food engineering': 'Agricultural Process Engineering',
    'irrigation and drainage engineering': 'Soil and Water Conservation Engineering'
}

SECTION_ORDER = [
    'General Aptitude',
    'Engineering Mathematics',
    'Farm Power and Machinery',
    'Soil and Water Conservation Engineering',
    'Agricultural Process Engineering'
]

def get_section_sort_key(sec):
    try:
        return SECTION_ORDER.index(sec)
    except ValueError:
        return 99

def clean_section(sec_raw):
    if not sec_raw:
        return 'General Aptitude'
    cleaned = sec_raw.strip().lower()
    cleaned = re.sub(r'^(section\s*\d*\s*[:—\-]?\s*)', '', cleaned).strip()
    for k, v in SECTION_MAP.items():
        if k in cleaned:
            return v
    if 'math' in cleaned:
        return 'Engineering Mathematics'
    if 'farm' in cleaned:
        return 'Farm Power and Machinery'
    if 'soil' in cleaned or 'water' in cleaned or 'drainage' in cleaned or 'irrigation' in cleaned:
        return 'Soil and Water Conservation Engineering'
    if 'process' in cleaned or 'food' in cleaned or 'agri' in cleaned or 'dairy' in cleaned:
        return 'Agricultural Process Engineering'
    return sec_raw.strip() or 'General Aptitude'

def clean_answer_key(raw_ans):
    if not raw_ans:
        return ""
    txt = raw_ans.strip()
    txt = re.sub(r'^(Official\s*(Key|Answer)\s*:?|Correct option\s*:?|Accepted range\s*:?|Key range\s*:?)', '', txt, flags=re.IGNORECASE).strip()
    
    m_range = re.search(r'(\d+\.?\d*\s*to\s*\d+\.?\d*)', txt, re.IGNORECASE)
    if m_range:
        return m_range.group(1).strip()
        
    opts = re.findall(r'\b([A-D])\b', txt)
    if opts:
        u_opts = []
        for o in opts:
            if o not in u_opts:
                u_opts.append(o)
        return ", ".join(u_opts)

    m_num = re.search(r'[-+]?\d*\.?\d+', txt)
    if m_num:
        return m_num.group(0)

    return txt

def extract_doc_images(docx_path, year, doc):
    """Extract embedded docx media images and map to paragraph indices"""
    p_img_map = {}
    with zipfile.ZipFile(docx_path) as z:
        rels_file = 'word/_rels/document.xml.rels'
        if rels_file not in z.namelist():
            return p_img_map
        rels_tree = ET.fromstring(z.read(rels_file))
        rel_map = {}
        for elem in rels_tree:
            rId = elem.attrib.get('Id')
            target = elem.attrib.get('Target')
            if rId and target:
                rel_map[rId] = target

        for p_idx, p in enumerate(doc.paragraphs):
            p_xml_str = p._p.xml
            if 'w:drawing' in p_xml_str or 'graphicData' in p_xml_str:
                p_xml = ET.fromstring(p_xml_str)
                embed_ids = []
                for elem in p_xml.iter():
                    for k, v in elem.attrib.items():
                        if k.endswith('embed'):
                            embed_ids.append(v)
                for rId in embed_ids:
                    if rId in rel_map:
                        img_target = rel_map[rId]
                        if img_target.startswith('media/'):
                            img_target = 'word/' + img_target
                        if img_target in z.namelist():
                            img_data = z.read(img_target)
                            ext = os.path.splitext(img_target)[1] or '.png'
                            out_name = f"gate_{year}_p{p_idx}{ext}"
                            out_path = os.path.join(OUT_IMG_DIR, out_name)
                            with open(out_path, 'wb') as f:
                                f.write(img_data)
                            p_img_map[p_idx] = f"/docx_images/{out_name}"
    return p_img_map

def parse_docx_file(docx_path, year):
    doc = docx.Document(docx_path)
    p_img_map = extract_doc_images(docx_path, year, doc)

    flat_lines = []
    for p_idx, p in enumerate(doc.paragraphs):
        p_text = p.text.strip()
        img_url = p_img_map.get(p_idx, None)
        if not p_text and not img_url:
            continue
        sub_lines = p_text.split('\n') if p_text else ['']
        for sl in sub_lines:
            flat_lines.append({
                'text': sl.strip(),
                'p_idx': p_idx,
                'img_url': img_url
            })

    q_blocks = []
    curr_block = []
    q_header_re = re.compile(r'^(?:Q|GA\s*Q|AG\s*Q|Question\s*GA-|Question\s*AG-|Question\s*|QAG|QGA)\s*[\.\-]?\s*(\d+)\b', re.IGNORECASE)

    for line_obj in flat_lines:
        txt = line_obj['text']
        is_header = False
        
        if not re.match(r'^(Question Statement|Question statement|Question Paper|Question Analysis|Question:)', txt, re.IGNORECASE):
            m = q_header_re.match(txt)
            if m:
                if len(txt) < 45 or ('[' in txt and 'mark' in txt) or ('(' in txt and 'mark' in txt):
                    is_header = True
        if txt.startswith('Section:'):
            curr_txt = ' '.join([lo['text'] for lo in curr_block])
            has_exp = re.search(r'(Detailed\s*Explanation|Explanation|Solution)', curr_txt, re.IGNORECASE) is not None
            if len(curr_block) == 0 or has_exp:
                is_header = True

        if is_header:
            if curr_block:
                q_blocks.append(curr_block)
            curr_block = [line_obj]
        else:
            if curr_block:
                curr_block.append(line_obj)
    if curr_block:
        q_blocks.append(curr_block)

    valid_q_blocks = []
    has_ans_line_re = re.compile(r'^(Official\s*(Answer|Key)|Answer\s*(Key|Key/Solution|\(.*?\)|:)|Accepted\s*Range|Correct\s*(option|answer)|Independently\s*(Derived|Computed)\s*Answer|Computed\s*Answer|Verified\s*Answer)', re.IGNORECASE)
    has_ans_block_re = re.compile(r'(Official\s*(Answer|Key)|Answer\s*(Key|Key/Solution|\(.*?\)|:)|Accepted\s*Range|Correct\s*(option|answer)|Independently\s*(Derived|Computed)\s*Answer|Computed\s*Answer|Verified\s*Answer|\([A-D]\))', re.IGNORECASE)

    for b in q_blocks:
        if not b:
            continue
        first_txt = b[0]['text'] if b else ''
        if not (q_header_re.match(first_txt) or first_txt.startswith('Section:')):
            continue
        if any(kw in first_txt for kw in ['Appendix', 'APPENDIX', 'Document contents', 'Fully Solved', 'Question Paper:', 'GATE 20', 'Prepared for:']):
            continue
        b_text = " ".join([lo['text'] for lo in b])
        if 'APPENDIX' in b_text[:40] or 'Appendix' in b_text[:40] or 'Document contents' in b_text[:40]:
            continue
        has_opts = re.search(r'\([A-D]\)', b_text) is not None or 'Options:' in b_text
        has_ans = has_ans_block_re.search(b_text) is not None
        has_exp = re.search(r'(Detailed\s*Explanation|Explanation|Solution)', b_text, re.IGNORECASE) is not None
        if (has_opts or has_ans) and has_exp:
            valid_q_blocks.append(b)

    parsed_questions = []

    for idx, block in enumerate(valid_q_blocks):
        block_text_all = "\n".join([lo['text'] for lo in block])
        
        qnum = idx + 1
        for lo in block[:3]:
            txt = lo['text']
            m = q_header_re.match(txt)
            if m and not re.match(r'^(Question Statement|Question statement|Question Paper|Question Analysis|Question:)', txt, re.IGNORECASE):
                num = int(m.group(1))
                if ('AG' in txt or 'QAG' in txt) and int(year) >= 2010 and num <= 55:
                    num += 10
                qnum = num
                break
            m2 = re.search(r'\(GA-(\d+)\)', txt, re.IGNORECASE)
            if m2:
                qnum = int(m2.group(1))
                break
            m3 = re.search(r'\(AG-(\d+)\)', txt, re.IGNORECASE)
            if m3:
                num = int(m3.group(1))
                if int(year) >= 2010 and num <= 55:
                    num += 10
                qnum = num
                break

        section = "General Aptitude"
        topic = "General Syllabus"
        subtopic = ""
        type_str = "MCQ"
        marks = 1
        block_img_url = ""

        for line_obj in block:
            if line_obj['img_url']:
                block_img_url = line_obj['img_url']
                break

        q_lines = []
        option_dict = {}
        ans_lines = []
        exp_lines = []

        mode = 'meta'

        for line_obj in block:
            line = line_obj['text']
            if not line and not line_obj['img_url']:
                continue

            if line.startswith('Section:'):
                sec_val = line.split('Section:', 1)[1].strip()
                sec_val = re.sub(r'Topic.*', '', sec_val).strip()
                section = clean_section(sec_val)
                
            if 'Topic' in line and (line.startswith('Topic') or 'Topic (' in line or 'Topic:' in line):
                top_val = re.sub(r'.*Topic(?:\s*\(.*?\))?\s*:\s*', '', line, flags=re.IGNORECASE).strip()
                top_val = re.sub(r'\[Note:.*\]', '', top_val).strip()
                if not top_val:
                    # check next line in block
                    l_idx = block.index(line_obj)
                    if l_idx + 1 < len(block):
                        next_l = block[l_idx + 1]['text'].strip()
                        if next_l and not re.match(r'^(Type|Question|Section|\[MCQ|\[MSQ|\[NAT)\b', next_l, re.IGNORECASE):
                            top_val = next_l
                if top_val:
                    if '—' in top_val:
                        t_parts = top_val.split('—', 1)
                        topic = t_parts[0].strip()
                        subtopic = t_parts[1].strip()
                    elif ':' in top_val and not top_val.startswith('http'):
                        t_parts = top_val.split(':', 1)
                        topic = t_parts[0].strip()
                        subtopic = t_parts[1].strip()
                    elif '-' in top_val:
                        t_parts = top_val.split('-', 1)
                        topic = t_parts[0].strip()
                        subtopic = t_parts[1].strip()
                    else:
                        topic = top_val

            if line.startswith('Type:') or line.startswith('Type / Marks:') or 'Type of question:' in line:
                if 'MSQ' in line: type_str = 'MSQ'
                elif 'NAT' in line: type_str = 'NAT'
                else: type_str = 'MCQ'
                if '2' in line: marks = 2
                continue

            if '[MCQ' in line or '[MSQ' in line or '[NAT' in line:
                if 'MSQ' in line: type_str = 'MSQ'
                elif 'NAT' in line: type_str = 'NAT'
                else: type_str = 'MCQ'
                if '2 mark' in line: marks = 2
                continue

            if re.match(r'^(Question|Question Statement)\s*:?', line, re.IGNORECASE):
                mode = 'question'
                rest = re.sub(r'^(Question|Question Statement)\s*:?', '', line, flags=re.IGNORECASE).strip()
                if rest and not re.match(r'^(Section|Topic|Type|Statement|Question Statement|Q\s*\.?\s*\d+|GA-\d+|AG-\d+|\[MCQ|\[MSQ|\[NAT)\b', rest, re.IGNORECASE):
                    q_lines.append(rest)
                continue

            if has_ans_line_re.match(line):
                mode = 'answer'
                ans_lines.append(line)
                continue

            if re.match(r'^(Detailed\s*Explanation|Explanation|Solution)', line, re.IGNORECASE):
                mode = 'solution'
                continue

            if mode == 'answer':
                ans_lines.append(line)
                continue

            if mode == 'solution':
                exp_lines.append(line)
                continue

            opt_matches = list(re.finditer(r'\(([A-D])\)\s*([^(\n]+)', line))
            if len(opt_matches) >= 2:
                for om in opt_matches:
                    option_dict[om.group(1).upper()] = om.group(2).strip()
                continue
            elif re.match(r'^\(([A-D])\)\s*(.*)', line):
                m_opt = re.match(r'^\(([A-D])\)\s*(.*)', line)
                option_dict[m_opt.group(1).upper()] = m_opt.group(2).strip()
                continue

            if mode in ['meta', 'question']:
                if not re.match(r'^(Section|Topic|Type|Type / Marks|Type of question|Question Statement|Statement|Figure \(reconstruction\)|Note|Q\s*\.?\s*\d+|GA-\d+|AG-\d+)\b', line, re.IGNORECASE):
                    q_lines.append(line)

        if qnum <= 10 and year != 2021:
            section = "General Aptitude"
        elif section == "General Aptitude" and qnum > 10:
            section = "Farm Power and Machinery"

        if not topic or not str(topic).strip():
            topic = f"{section} Core Concepts"

        question_text = "\n".join(q_lines).strip()
        question_text = re.sub(r'^(Question|Question Statement|Statement)\s*:?\s*', '', question_text, flags=re.IGNORECASE).strip()
        question_text = re.sub(r'Topic\s*\(syllabus\)\s*:.*?\n', '', question_text, flags=re.IGNORECASE).strip()

        solution_text = "\n".join(exp_lines).strip()
        ans_raw = " ".join(ans_lines).strip()
        correct_answer = clean_answer_key(ans_raw)

        if not correct_answer:
            m_range = re.search(r'Official\s*Accepted\s*Range\s*:\s*(\d+\.?\d*\s*to\s*\d+\.?\d*)', question_text, re.IGNORECASE)
            if m_range:
                correct_answer = m_range.group(1).strip()
            else:
                m_num = re.search(r'Precisely\s*Computed\s*Value\s*:\s*([-+]?\d*\.?\d+)', question_text, re.IGNORECASE)
                if m_num:
                    correct_answer = m_num.group(1).strip()

        if not solution_text:
            solution_text = f"Official Verified Key: {correct_answer}"

        q_obj = {
            "id": f"GATE_{year}_Q{qnum}",
            "year": str(year),
            "qnum": qnum,
            "section": section,
            "topic": topic,
            "subtopic": subtopic,
            "gate_section": "GA" if section == "General Aptitude" else "AG",
            "type": type_str,
            "marks": marks,
            "negative_marks": (1/3 if marks == 1 else 2/3) if type_str == "MCQ" else 0,
            "question": question_text,
            "options": option_dict,
            "correct_answer": correct_answer,
            "solution": solution_text,
            "snippet_url": "",
            "image_url": block_img_url
        }

        parsed_questions.append(q_obj)

    return parsed_questions

def main():
    all_qs = []
    paper_qs_map = {}

    solved_files = sorted([f for f in os.listdir(SOLVED_DIR) if f.endswith('.docx')])
    print(f"Found {len(solved_files)} solved DOCX files: {solved_files}")

    for fn in solved_files:
        m = re.search(r'(\d{4})', fn)
        if not m:
            continue
        year = int(m.group(1))
        fp = os.path.join(SOLVED_DIR, fn)
        qs = parse_docx_file(fp, year)
        print(f"Extracted {len(qs)} questions for GATE {year} from {fn}.")
        all_qs.extend(qs)
        paper_qs_map[str(year)] = qs

    print(f"Total extracted questions from solved DOCX files: {len(all_qs)}")

    # Sort questions according to Section, Topic, Subtopic, Year, and Qnum
    all_qs.sort(key=lambda q: (
        get_section_sort_key(q['section']),
        q['topic'].lower(),
        q['subtopic'].lower(),
        int(q['year']),
        q['qnum']
    ))

    with open(OUT_JSON, 'w', encoding='utf-8') as f:
        json.dump(all_qs, f, indent=2, ensure_ascii=False)

    print(f"Successfully generated clean, topic-sorted practice dataset in {OUT_JSON}")

    # Now build mock_papers.json for all 20 years (2007-2026) strictly based on DOCX availability
    mock_papers = []
    all_years_list = [str(y) for y in range(2026, 2006, -1)]

    for year_str in all_years_list:
        y_num = int(year_str)
        is_solved_avail = year_str in paper_qs_map and len(paper_qs_map[year_str]) > 0

        max_marks = 150 if y_num <= 2009 else 100
        total_qs = 85 if y_num <= 2009 else 65
        ga_qs = 0 if y_num <= 2009 else 10
        ag_qs = total_qs - ga_qs

        p_qs = paper_qs_map.get(year_str, [])

        if is_solved_avail:
            instructions = [
                f"1. Total duration of GATE {year_str} Examination is 180 minutes (3 Hours).",
                f"2. The paper contains {len(p_qs)} verified questions carrying {max_marks} Maximum Marks.",
                "3. Section 1 (Q.1-Q.10 GA) carries 15 marks. Section 2 (AG) carries 85 marks.",
                "4. Negative marking: 1/3 mark for 1-mark MCQs, 2/3 mark for 2-mark MCQs. 0 for MSQ/NAT.",
                "5. Official verified solutions and derivations provided for every question."
            ]
        else:
            instructions = [
                f"Detailed Solved .docx Paper for GATE {year_str} is currently being verified and formatted.",
                "Adding Soon!!!"
            ]

        mock_papers.append({
            "year": year_str,
            "title": f"GATE {year_str} Agricultural Engineering Paper",
            "status": "AVAILABLE" if is_solved_avail else "ADDING_SOON",
            "has_solved_docx": is_solved_avail,
            "instructions": {
                "year": year_str,
                "duration_mins": 180,
                "max_marks": max_marks,
                "total_qs": total_qs,
                "ga_qs": ga_qs,
                "ag_qs": ag_qs,
                "is_untimed": False,
                "instructions": instructions
            },
            "questions": p_qs
        })

    mock_out_path = 'src/data/mock_papers.json'
    with open(mock_out_path, 'w', encoding='utf-8') as f:
        json.dump(mock_papers, f, indent=2, ensure_ascii=False)

    print(f"Successfully generated {len(mock_papers)} paper entries in {mock_out_path}")

if __name__ == '__main__':
    main()
